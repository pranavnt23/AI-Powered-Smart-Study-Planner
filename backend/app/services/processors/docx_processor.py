from pathlib import Path
from docx import Document

class DOCXProcessor:
    @staticmethod
    def extract_text(file_path: str) -> dict:
        try:
            document = Document(file_path)
            extracted_lines = []

            def append_text(text: str):
                if not text:
                    return
                cleaned = text.strip()
                if cleaned:
                    extracted_lines.append(cleaned)

            for paragraph in document.paragraphs:
                append_text(paragraph.text)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        append_text(cell.text)

            for section in document.sections:
                header = section.header
                footer = section.footer

                for paragraph in header.paragraphs:
                    append_text(paragraph.text)

                for paragraph in footer.paragraphs:
                    append_text(paragraph.text)

            extracted_text = "\n".join(extracted_lines)

            return {
                "status": True,
                "file_type": "docx",
                "file_name": Path(file_path).name,
                "content": extracted_text,
            }
        except Exception as error:
            return {
                "status": False,
                "file_type": "docx",
                "file_name": Path(file_path).name,
                "content": "",
                "error": str(error),
            }